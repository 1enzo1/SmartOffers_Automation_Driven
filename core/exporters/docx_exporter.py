import json
from io import BytesIO

from docx import Document
from docx.shared import Pt

from .common import render_planned_content
from .storage import persist_export_bytes


def _set_normal_style(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)


def _add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def _add_kv_table(document, title, items):
    _add_heading(document, title, level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Campo"
    headers[1].text = "Valor"

    for key, value in items:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = "" if value is None else str(value)


def _add_list(document, title, items):
    _add_heading(document, title, level=2)

    if not items:
        document.add_paragraph("Nenhum item registrado.")
        return

    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _add_json_block(document, title, data):
    _add_heading(document, title, level=2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(json.dumps(data, indent=2, ensure_ascii=False))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def _build_document(context):
    document = Document()
    _set_normal_style(document)

    document.add_heading(context["title"], level=0)
    if context.get("summary"):
        document.add_paragraph(context["summary"])

    _add_kv_table(
        document,
        "Metadados da exportacao",
        [
            ("generated_at", context["metadata"]["generated_at"]),
            ("export_type", context["metadata"]["export_type"]),
            ("source_type", context["metadata"]["source_type"]),
            ("source_id", context["metadata"]["source_id"]),
        ],
    )

    scenario = context.get("scenario") or {}
    answers = scenario.get("source_answers") or {}
    if scenario:
        _add_kv_table(
            document,
            "Resumo do cenario",
            [
                ("cenario_id", scenario.get("id", "")),
                ("titulo", scenario.get("titulo", "")),
                ("campaign_id", answers.get("campaign_id", "")),
                ("campaign_name", answers.get("campaign_name", "")),
                ("customer_type", answers.get("customer_type", "")),
                ("document_type", answers.get("document_type", "")),
                ("event_type", answers.get("event_type", "")),
                ("objective", answers.get("objective", "")),
            ],
        )

    _add_kv_table(
        document,
        "Resumo de validacao",
        [
            ("execution_steps", len(context.get("execution_steps") or [])),
            ("validation_steps", len(context.get("validation_steps") or [])),
            ("queries", len(context.get("queries") or [])),
            ("checkpoints", len(context.get("checkpoints") or [])),
            ("evidence_files", len(context.get("evidence_files") or [])),
        ],
    )

    _add_json_block(document, "Payload", context.get("payload") or {})

    execution_steps = context.get("execution_steps") or []
    if execution_steps:
        _add_heading(document, "Execution Steps", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Step"
        headers[1].text = "Acao"
        headers[2].text = "Detalhes"
        headers[3].text = "Esperado"
        for item in execution_steps:
            cells = table.add_row().cells
            cells[0].text = str(item.get("step", ""))
            cells[1].text = str(item.get("action", ""))
            cells[2].text = str(item.get("details", ""))
            cells[3].text = str(item.get("expected_result", ""))

    validation_steps = context.get("validation_steps") or []
    if validation_steps:
        _add_heading(document, "Validation Steps", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Step"
        headers[1].text = "Validacao"
        headers[2].text = "Detalhes"
        headers[3].text = "Esperado"
        for item in validation_steps:
            cells = table.add_row().cells
            cells[0].text = str(item.get("step", ""))
            cells[1].text = str(item.get("validation", ""))
            cells[2].text = str(item.get("details", ""))
            cells[3].text = str(item.get("expected_result", ""))

    if context.get("queries"):
        _add_heading(document, "Queries", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = table.rows[0].cells
        headers[0].text = "Nome"
        headers[1].text = "Tipo"
        headers[2].text = "Objetivo"
        headers[3].text = "Consulta"
        for item in context["queries"]:
            cells = table.add_row().cells
            cells[0].text = str(item.get("name", ""))
            cells[1].text = str(item.get("kind", ""))
            cells[2].text = str(item.get("purpose", ""))
            cells[3].text = render_planned_content(item)

    _add_list(document, "Checkpoints", context.get("checkpoints") or [])
    _add_list(document, "Evidencias esperadas", context.get("evidence_files") or [])
    _add_list(document, "Avisos", context.get("warnings") or [])

    dry_run = context.get("dry_run") or {}
    if dry_run:
        _add_kv_table(
            document,
            "Dry-run Results",
            [
                ("report_id", dry_run.get("id", "")),
                ("scenario_id", dry_run.get("scenario_id", "")),
                ("status", dry_run.get("status", "")),
                ("started_at", dry_run.get("started_at", "")),
                ("finished_at", dry_run.get("finished_at", "")),
                ("duration_ms", dry_run.get("duration_ms", "")),
                ("total_steps", (dry_run.get("summary") or {}).get("total", "")),
                ("passed", (dry_run.get("summary") or {}).get("passed", "")),
                ("failed", (dry_run.get("summary") or {}).get("failed", "")),
                ("skipped", (dry_run.get("summary") or {}).get("skipped", "")),
            ],
        )

        dry_run_steps = dry_run.get("steps") or []
        if dry_run_steps:
            _add_heading(document, "Dry-run Steps", level=2)
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            headers = table.rows[0].cells
            headers[0].text = "Tipo"
            headers[1].text = "Nome"
            headers[2].text = "Status"
            headers[3].text = "Duracao (ms)"
            headers[4].text = "Mensagem"
            for item in dry_run_steps:
                cells = table.add_row().cells
                cells[0].text = str(item.get("type", ""))
                cells[1].text = str(item.get("name", ""))
                cells[2].text = str(item.get("status", ""))
                cells[3].text = str(item.get("duration_ms", ""))
                cells[4].text = str(item.get("message", ""))

        _add_list(document, "Logs", dry_run.get("logs") or [])

    return document


def export_docx_artifact(context):
    document = _build_document(context)
    buffer = BytesIO()
    document.save(buffer)
    return persist_export_bytes(
        buffer.getvalue(),
        context["source_type"],
        context["source_id"],
        context["metadata"]["export_type"],
        "docx",
    )
