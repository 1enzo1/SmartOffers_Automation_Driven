import json
import uuid
from pathlib import Path

import app as app_module
from docx import Document
from openpyxl import load_workbook


ROOT = Path(".test_output") / "exports"


def valid_payload(**overrides):
    payload = {
        "campaign_name": "Squad162 Upsell",
        "campaign_id": "162",
        "system": "SmartOffers",
        "objective": "Validar bonificacao apenas para upgrade",
        "customer_type": "pos",
        "document_type": "PF",
        "event_type": "upsell",
        "validations": ["api", "database", "audit"],
        "deadline_rule": "d1",
    }
    payload.update(overrides)
    return payload


def make_client(monkeypatch):
    base = ROOT / uuid.uuid4().hex
    monkeypatch.setenv("CENARIOS_GERADOS_PATH", str(base / "cenarios"))
    monkeypatch.setenv("DRYRUNS_GERADOS_PATH", str(base / "dryruns"))
    monkeypatch.setenv("EXPORTS_GERADOS_PATH", str(base / "exports"))
    app_module.app.config.update(TESTING=True)
    return app_module.app.test_client(), base


def latest_file(directory, suffix):
    files = sorted(directory.glob(f"*.{suffix}"), key=lambda path: path.stat().st_mtime)
    assert files, f"no files with suffix {suffix} in {directory}"
    return files[-1]


def docx_text(path):
    document = Document(str(path))
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_exports_cover_scenario_and_dry_run(monkeypatch):
    client, base = make_client(monkeypatch)

    scenario_response = client.post("/api/scenarios/generate", json=valid_payload())
    assert scenario_response.status_code == 201
    scenario = scenario_response.get_json()["scenario"]

    dry_run_response = client.post(f"/api/scenarios/{scenario['id']}/dry-run")
    assert dry_run_response.status_code == 201
    report = dry_run_response.get_json()["report"]

    exports_dir = base / "exports"

    scenario_json_response = client.get(f"/api/scenarios/{scenario['id']}/export/json")
    assert scenario_json_response.status_code == 200
    scenario_json_path = latest_file(exports_dir, "json")
    scenario_json = json.loads(scenario_json_path.read_text(encoding="utf-8"))
    assert scenario_json["metadata"]["export_type"] == "json"
    assert scenario_json["metadata"]["source_type"] == "scenario"
    assert scenario_json["metadata"]["source_id"] == scenario["id"]
    assert "known_limitations" in scenario_json["metadata"]
    assert scenario_json["scenario"]["id"] == scenario["id"]

    scenario_docx_response = client.get(f"/api/scenarios/{scenario['id']}/export/docx")
    assert scenario_docx_response.status_code == 200
    scenario_docx_path = latest_file(exports_dir, "docx")
    scenario_docx_content = docx_text(scenario_docx_path)
    assert scenario["titulo"] in scenario_docx_content
    assert "Execution Steps" in scenario_docx_content
    assert "Validation Steps" in scenario_docx_content
    assert "Payload" in scenario_docx_content

    scenario_xlsx_response = client.get(f"/api/scenarios/{scenario['id']}/export/xlsx")
    assert scenario_xlsx_response.status_code == 200
    scenario_xlsx_path = latest_file(exports_dir, "xlsx")
    scenario_wb = load_workbook(scenario_xlsx_path)
    assert scenario_wb.sheetnames == [
        "Resumo",
        "Execution Steps",
        "Validation Steps",
        "Payload",
        "Queries",
        "Checkpoints",
        "Evidencias",
        "Warnings",
    ]

    report_json_response = client.get(f"/api/dry-runs/{report['id']}/export/json")
    assert report_json_response.status_code == 200
    report_json_path = latest_file(exports_dir, "json")
    report_json = json.loads(report_json_path.read_text(encoding="utf-8"))
    assert report_json["metadata"]["export_type"] == "json"
    assert report_json["metadata"]["source_type"] == "dry_run"
    assert report_json["metadata"]["source_id"] == report["id"]
    assert report_json["dry_run"]["id"] == report["id"]
    assert report_json["scenario"]["id"] == scenario["id"]

    report_docx_response = client.get(f"/api/dry-runs/{report['id']}/export/docx")
    assert report_docx_response.status_code == 200
    report_docx_path = latest_file(exports_dir, "docx")
    report_docx_content = docx_text(report_docx_path)
    assert report["id"] in report_docx_content
    assert "Dry-run Results" in report_docx_content
    assert "Dry-run Steps" in report_docx_content
    assert "Logs" in report_docx_content

    report_xlsx_response = client.get(f"/api/dry-runs/{report['id']}/export/xlsx")
    assert report_xlsx_response.status_code == 200
    report_xlsx_path = latest_file(exports_dir, "xlsx")
    report_wb = load_workbook(report_xlsx_path)
    assert report_wb.sheetnames == [
        "Resumo",
        "Execution Steps",
        "Validation Steps",
        "Payload",
        "Queries",
        "Checkpoints",
        "Evidencias",
        "Warnings",
        "Dry-run Results",
        "Dry-run Logs",
    ]


def test_export_and_lookup_404s(monkeypatch):
    client, _ = make_client(monkeypatch)

    assert client.get("/api/scenarios/nao-existe/export/docx").status_code == 404
    assert client.get("/api/dry-runs/nao-existe/export/json").status_code == 404
    assert client.get("/api/dry-runs/nao-existe").status_code == 404
    assert client.post("/api/scenarios/nao-existe/dry-run").status_code == 404
